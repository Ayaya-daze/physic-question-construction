"""DOCX parser — extracts text, detects formulas, and extracts embedded images."""

from __future__ import annotations

from pathlib import Path
from typing import Optional

from app.services.parsers import (
    MediaRef,
    PageContent,
    ParsedDocument,
    TextBlock,
    register_parser,
)


def _detect_formula(paragraph) -> Optional[str]:
    """Try to detect OMML math objects in a paragraph and convert to LaTeX.

    Returns LaTeX string if found, None otherwise.
    """
    try:
        # Check for Office Math Markup Language (OMML) elements
        # python-docx doesn't directly support OMML parsing, but we can detect it
        element = paragraph._element
        # Look for m:oMath or m:oMathPara namespaced elements
        math_ns = "http://schemas.openxmlformats.org/officeDocument/2006/math"
        math_elements = element.findall(f"{{{math_ns}}}oMath")
        math_para = element.findall(f"{{{math_ns}}}oMathPara")

        if math_elements or math_para:
            # Extract raw math text from m:t elements
            math_text_parts = []
            for math_el in math_elements + math_para:
                texts = math_el.findall(f".//{{{math_ns}}}t")
                math_text_parts.extend(t.get("value", t.text or "") for t in texts)
            math_text = " ".join(math_text_parts).strip()
            if math_text:
                return f"${math_text}$"
    except Exception:
        pass

    return None


def _find_embedded_images(paragraph, page_num: int = 1) -> list[MediaRef]:
    """Extract references to embedded images in a paragraph."""
    refs: list[MediaRef] = []
    try:
        from docx.opc.constants import RELATIONSHIP_TYPE as RT

        element = paragraph._element
        nsmap = {
            "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
            "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
            "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
        }

        drawings = element.findall(".//wp:inline", nsmap)
        for i, _drawing in enumerate(drawings):
            refs.append(MediaRef(
                asset_id=f"asset_img_{i}",
                asset_type="figure",
                file_path="",
                page_number=page_num,
            ))
    except Exception:
        pass

    return refs


@register_parser("docx")
def parse_docx(file_path: Path) -> ParsedDocument:
    """Parse a DOCX file, extracting text with paragraph styles and detecting formulas."""
    from docx import Document

    doc = Document(str(file_path))
    blocks: list[TextBlock] = []
    raw_text_lines: list[str] = []
    all_media_refs: list[MediaRef] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        style_name = para.style.name if para.style else "Normal"
        block_type = "paragraph"
        if style_name and any(h in style_name.lower() for h in ("heading", "title", "标题")):
            block_type = "heading"

        # Check for formulas
        formula_text = _detect_formula(para)
        if formula_text:
            blocks.append(TextBlock(text=formula_text, block_type="formula"))
            raw_text_lines.append(formula_text)
        else:
            blocks.append(TextBlock(text=text, block_type=block_type))
            raw_text_lines.append(text)

        # Check for images
        media_refs = _find_embedded_images(para)
        all_media_refs.extend(media_refs)

    # Extract document metadata
    metadata: dict = {}
    try:
        props = doc.core_properties
        metadata["title"] = props.title or None
        metadata["author"] = props.author or None
    except Exception:
        pass

    raw_text = "\n".join(raw_text_lines)

    return ParsedDocument(
        pages=[PageContent(page_number=1, blocks=blocks, raw_text=raw_text)],
        raw_text=raw_text,
        metadata={**metadata, "format": "docx", "paragraph_count": len(doc.paragraphs)},
        media_refs=all_media_refs,
    )
